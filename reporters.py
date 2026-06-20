# reporters.py
#
# Node 7 — Analytics Agent: output reporting
#
# Two structured outputs are generated at the end of each pipeline run:
#
#   1. scores.csv  — one row per student, all scores + similarity + error
#   2. feedback.docx — formatted per-student feedback report with:
#        - AI score breakdown (total + per criterion)
#        - Human score breakdown (total + per criterion)
#        - AI-generated qualitative feedback
#        - Human-authored qualitative feedback
#        - Semantic similarity score
#        - Score error (AI − human)

import csv
import os
from datetime import datetime

from config import CRITERIA, OUTPUT_CSV, OUTPUT_FEEDBACK, RESULTS_DIR


# ══════════════════════════════════════════════════════════════════════════════
# 📊  CSV EXPORT  —  scores.csv
# ══════════════════════════════════════════════════════════════════════════════

def export_scores_csv(results: dict, correlations: dict = None,
                      errors: list = None, feedback_sims: list = None):
    """
    Write a CSV file with one row per student containing:
      - Student ID
      - AI total score
      - AI criterion scores (one column each)
      - Human total score
      - Human criterion scores (one column each)
      - Score error (AI − human)
      - Semantic similarity
      - TF-IDF similarity (if available)

    A summary row is appended at the bottom with cohort-level statistics.

    Parameters
    ----------
    results      : dict — the full results dict from main.py
    correlations : dict — {"pearson": float, "spearman": float} or None
    errors       : list of floats — signed score errors
    feedback_sims: list of floats — semantic similarity scores
    """
    os.makedirs(RESULTS_DIR, exist_ok=True)

    criteria_keys = list(CRITERIA.keys())

    # Build column headers
    ai_crit_cols    = [f"AI_{c.capitalize()}" for c in criteria_keys]
    human_crit_cols = [f"Human_{c.capitalize()}" for c in criteria_keys]

    headers = (
        ["Student_ID", "AI_Total", "Human_Total", "Error_AI_minus_Human",
         "Semantic_Similarity", "Review_Status", "Decision", "Override_Score",
         "Flagged", "Flag_Reasons", "Reviewer_Note"]
        + ai_crit_cols
        + human_crit_cols
    )

    rows = []
    for student_id, data in results.items():
        ai     = data.get("ai", {})
        human  = data.get("human", {})
        review = data.get("review", {})

        ai_criteria    = ai.get("criteria_scores", {}) or {}
        human_criteria = human.get("criteria_scores", {}) or {}

        reviewed = review.get("reviewed", False)
        flagged  = review.get("flagged", False)
        decision = review.get("decision", "")
        reasons  = " | ".join(review.get("reasons", []))

        row = {
            "Student_ID":            _clean_id(student_id),
            "AI_Total":              ai.get("total_score", ""),
            "Human_Total":           human.get("total_score", ""),
            "Error_AI_minus_Human":  data.get("error", ""),
            "Semantic_Similarity":   data.get("similarity", ""),
            "Review_Status":         "Reviewed" if reviewed else "Pending",
            "Decision":              decision or "",
            "Override_Score":        review.get("override_score", "") or "",
            "Flagged":               "Yes" if flagged else "No",
            "Flag_Reasons":          reasons,
            "Reviewer_Note":         review.get("reviewer_note", "") or "",
        }

        for c in criteria_keys:
            row[f"AI_{c.capitalize()}"]    = ai_criteria.get(c, "")
            row[f"Human_{c.capitalize()}"] = human_criteria.get(c, "")

        rows.append(row)

    with open(OUTPUT_CSV, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=headers)
        writer.writeheader()
        writer.writerows(rows)

        # ── Summary section ────────────────────────────────────────────────
        f.write("\n")
        writer.writerow({h: "" for h in headers})  # blank separator row

        summary = {h: "" for h in headers}
        summary["Student_ID"] = "=== COHORT SUMMARY ==="
        writer.writerow(summary)

        if errors:
            import statistics
            err_row           = {h: "" for h in headers}
            err_row["Student_ID"]           = "Mean Error (AI − Human)"
            err_row["Error_AI_minus_Human"] = round(statistics.mean(errors), 3)
            writer.writerow(err_row)

            mae_row           = {h: "" for h in headers}
            mae_row["Student_ID"]           = "Mean Absolute Error"
            mae_row["Error_AI_minus_Human"] = round(
                sum(abs(e) for e in errors) / len(errors), 3
            )
            writer.writerow(mae_row)

        if feedback_sims:
            import statistics
            sim_row                   = {h: "" for h in headers}
            sim_row["Student_ID"]     = "Avg Semantic Similarity"
            sim_row["Semantic_Similarity"] = round(statistics.mean(feedback_sims), 3)
            writer.writerow(sim_row)

        if correlations:
            for label, key in [("Pearson r", "pearson"), ("Spearman rho", "spearman")]:
                corr_row              = {h: "" for h in headers}
                corr_row["Student_ID"] = label
                val = correlations.get(key)
                corr_row["AI_Total"]   = round(val, 4) if val is not None else "N/A"
                writer.writerow(corr_row)

        # Run metadata
        meta_row = {h: "" for h in headers}
        meta_row["Student_ID"] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        writer.writerow(meta_row)

    print(f"📊 Scores CSV saved → {OUTPUT_CSV}")


# ══════════════════════════════════════════════════════════════════════════════
# 📄  WORD FEEDBACK REPORT  —  feedback.docx
# ══════════════════════════════════════════════════════════════════════════════

def export_feedback_docx(results: dict, correlations: dict = None):
    """
    Generate a formatted Word document with detailed per-student feedback.

    Each student gets a section containing:
      - Score comparison table (AI vs human, total + per criterion)
      - AI-generated qualitative feedback
      - Human-authored qualitative feedback
      - Semantic similarity and error metrics

    A cohort summary page is appended at the end.

    Parameters
    ----------
    results      : dict — the full results dict from main.py
    correlations : dict — {"pearson": float, "spearman": float} or None
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("⚠️  python-docx not installed. Run: pip install python-docx")
        print("   Skipping feedback.docx export.")
        return

    os.makedirs(RESULTS_DIR, exist_ok=True)
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title page ────────────────────────────────────────────────────────
    _add_heading(doc, "AI-Assisted Essay Evaluation", level=0,
                 color=RGBColor(0x2E, 0x75, 0xB6))
    _add_heading(doc, "Detailed Feedback Report", level=1,
                 color=RGBColor(0x1F, 0x4E, 0x79))
    _add_para(doc,
              f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}\n"
              f"Total students evaluated: {len(results)}",
              size=10, color=RGBColor(0x80, 0x80, 0x80))
    doc.add_page_break()

    criteria_keys = list(CRITERIA.keys())

    # ── Per-student sections ──────────────────────────────────────────────
    for i, (student_id, data) in enumerate(results.items(), 1):
        ai    = data.get("ai", {})
        human = data.get("human", {})
        sim   = data.get("similarity", None)
        error = data.get("error", None)

        ai_criteria    = ai.get("criteria_scores", {}) or {}
        human_criteria = human.get("criteria_scores", {}) or {}
        ai_total       = ai.get("total_score", "N/A")
        human_total    = human.get("total_score", "N/A")

        # Student heading
        _add_heading(doc, f"Student {i}: {_clean_id(student_id)}",
                     level=1, color=RGBColor(0x2E, 0x75, 0xB6))

        # ── Score comparison table ─────────────────────────────────────
        _add_heading(doc, "Score Comparison", level=2,
                     color=RGBColor(0x1F, 0x4E, 0x79))

        col_widths = [Inches(2.0), Inches(1.5), Inches(1.5), Inches(1.5)]
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr = tbl.rows[0].cells
        for cell, text in zip(hdr, ["Criterion", "Max", "AI Score", "Human Score"]):
            _shade_cell(cell, "2E75B6")
            _set_cell_text(cell, text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

        # Criterion rows
        for c in criteria_keys:
            max_s = CRITERIA[c]
            ai_s  = ai_criteria.get(c, "—")
            hu_s  = human_criteria.get(c, "—")
            row   = tbl.add_row().cells
            _shade_cell(row[0], "F0F5FF")
            _set_cell_text(row[0], c.capitalize(), bold=True)
            _set_cell_text(row[1], str(max_s))
            _set_cell_text(row[2], str(ai_s),  color=_score_color(ai_s, max_s))
            _set_cell_text(row[3], str(hu_s))

        # Total row
        tot_row = tbl.add_row().cells
        _shade_cell(tot_row[0], "1F4E79")
        _shade_cell(tot_row[1], "1F4E79")
        _shade_cell(tot_row[2], "1F4E79")
        _shade_cell(tot_row[3], "1F4E79")
        _set_cell_text(tot_row[0], "TOTAL",       bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _set_cell_text(tot_row[1], "100",         bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _set_cell_text(tot_row[2], str(ai_total), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _set_cell_text(tot_row[3], str(human_total), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

        # Metrics below table
        doc.add_paragraph()
        metrics = []
        if error is not None:
            metrics.append(f"Score Error (AI − Human): {error:+.2f}")
        if sim is not None:
            metrics.append(f"Semantic Feedback Similarity: {sim:.3f}")
        if metrics:
            _add_para(doc, "  |  ".join(metrics), size=9,
                      color=RGBColor(0x44, 0x44, 0x44))

        # ── AI feedback ────────────────────────────────────────────────
        doc.add_paragraph()
        _add_heading(doc, "AI-Generated Feedback", level=2,
                     color=RGBColor(0x2E, 0x75, 0xB6))
        ai_feedback = ai.get("feedback", "No feedback generated.")
        _add_para(doc, ai_feedback or "No feedback generated.")

        # ── Human feedback ─────────────────────────────────────────────
        doc.add_paragraph()
        _add_heading(doc, "Human Assessor Feedback", level=2,
                     color=RGBColor(0x1F, 0x4E, 0x79))
        human_feedback = human.get("feedback", "No feedback available.")
        _add_para(doc, human_feedback or "No feedback available.",
                  color=RGBColor(0x33, 0x33, 0x33))

        # ── AI reasoning log (collapsed / smaller text) ────────────────
        # ---- Human review status ─────────────────────────────────────
        review = data.get("review", {})
        if review:
            doc.add_paragraph()
            _add_heading(doc, "Human Review Status", level=2,
                         color=RGBColor(0x2E, 0x75, 0xB6))

            flagged  = review.get("flagged", False)
            reviewed = review.get("reviewed", False)
            decision = review.get("decision", "") or ""
            reasons  = review.get("reasons", [])
            note     = review.get("reviewer_note", "") or ""
            override = review.get("override_score")

            if reviewed and decision == "overridden":
                status_text  = f"OVERRIDDEN -> {override}"
                status_color = RGBColor(0xB8, 0x6B, 0x00)
            elif reviewed and decision == "approved":
                status_text  = "APPROVED"
                status_color = RGBColor(0x1F, 0x7A, 0x1F)
            elif reviewed and decision == "flagged":
                status_text  = "FLAGGED BY REVIEWER"
                status_color = RGBColor(0xC0, 0x00, 0x00)
            elif flagged:
                status_text  = "FLAGGED - Pending Review"
                status_color = RGBColor(0xC0, 0x00, 0x00)
            else:
                status_text  = "Not Flagged"
                status_color = RGBColor(0x88, 0x88, 0x88)

            _add_para(doc, f"Status: {status_text}", bold=True, color=status_color)

            if reasons:
                _add_para(doc, "Automatic flag reasons:", size=10)
                for r in reasons:
                    _add_para(doc, f"  * {r}", size=10,
                              color=RGBColor(0xC0, 0x00, 0x00))

            if note:
                _add_para(doc, f"Reviewer note: {note}", size=10,
                          color=RGBColor(0x44, 0x44, 0x44))

        # ---- AI reasoning log ----------------------------------------
        reasoning_logs = ai.get("reasoning_logs", [])
        if reasoning_logs:
            doc.add_paragraph()
            _add_heading(doc, "AI Reasoning Log (Chain-of-Thought)", level=2,
                         color=RGBColor(0x70, 0x70, 0x70))
            for j, log in enumerate(reasoning_logs, 1):
                if log:
                    _add_para(doc, f"[Chunk {j}] {log}", size=8,
                              color=RGBColor(0x88, 0x88, 0x88))

        if i < len(results):
            doc.add_page_break()

    # ── Cohort summary page ───────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "Cohort Summary", level=0,
                 color=RGBColor(0x2E, 0x75, 0xB6))

    ai_totals    = [d["ai"]["total_score"]    for d in results.values()
                    if d.get("ai", {}).get("total_score") is not None]
    human_totals = [d["human"]["total_score"] for d in results.values()
                    if d.get("human", {}).get("total_score") is not None]
    sims         = [d["similarity"]           for d in results.values()
                    if d.get("similarity") is not None]
    errs         = [d["error"]                for d in results.values()
                    if d.get("error") is not None]

    import statistics as _stats

    summary_rows = []
    if ai_totals:
        summary_rows.append(("Mean AI Score",    f"{_stats.mean(ai_totals):.2f}"))
        summary_rows.append(("Mean Human Score", f"{_stats.mean(human_totals):.2f}"))
    if errs:
        summary_rows.append(("Mean Error (AI − Human)",
                              f"{_stats.mean(errs):+.2f}"))
        summary_rows.append(("Mean Absolute Error",
                              f"{sum(abs(e) for e in errs)/len(errs):.2f}"))
    if sims:
        summary_rows.append(("Avg Semantic Similarity",
                              f"{_stats.mean(sims):.3f}"))
    if correlations:
        p = correlations.get("pearson")
        s = correlations.get("spearman")
        summary_rows.append(("Pearson r",
                              f"{p:.4f}" if p is not None else "N/A"))
        summary_rows.append(("Spearman ρ",
                              f"{s:.4f}" if s is not None else "N/A"))

    if summary_rows:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        _shade_cell(hdr[0], "2E75B6")
        _shade_cell(hdr[1], "2E75B6")
        _set_cell_text(hdr[0], "Metric",    bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _set_cell_text(hdr[1], "Value",     bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        for label, value in summary_rows:
            row = tbl.add_row().cells
            _shade_cell(row[0], "F0F5FF")
            _set_cell_text(row[0], label, bold=True)
            _set_cell_text(row[1], value)

    doc.save(OUTPUT_FEEDBACK)
    print(f"📄 Feedback report saved → {OUTPUT_FEEDBACK}")


# ══════════════════════════════════════════════════════════════════════════════
# 🔧  HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def _clean_id(student_id):
    """Remove file extension from student ID for display."""
    return os.path.splitext(student_id)[0]


def _add_heading(doc, text, level=1, color=None):
    if level == 0:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = __import__('docx').shared.Pt(18)
        if color:
            run.font.color.rgb = color
        p.paragraph_format.space_after = __import__('docx').shared.Pt(6)
    else:
        h = doc.add_heading(text, level=level)
        if color:
            for run in h.runs:
                run.font.color.rgb = color


def _add_para(doc, text, size=11, color=None, bold=False):
    from docx.shared import Pt
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = __import__('docx').shared.Pt(4)
    return p


def _shade_cell(cell, hex_color):
    """Apply a background shade to a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)


def _set_cell_text(cell, text, bold=False, color=None, size=10):
    from docx.shared import Pt
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _score_color(score, max_score):
    """Return an RGBColor based on score percentage."""
    from docx.shared import RGBColor
    try:
        pct = float(score) / float(max_score)
    except (TypeError, ValueError, ZeroDivisionError):
        return None
    if pct >= 0.70:
        return RGBColor(0x1F, 0x7A, 0x1F)   # green — distinction
    elif pct >= 0.50:
        return RGBColor(0xB8, 0x6B, 0x00)   # amber — pass/merit
    else:
        return RGBColor(0xC0, 0x00, 0x00)   # red   — fail/bare pass